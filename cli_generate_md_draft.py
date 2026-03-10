import argparse
import json
import os
import re
from typing import List, Dict, Any, Optional, Callable
from src.img_agent.deepseek_client import DeepSeekClient
from src.config.env import get_env_config
from docx import Document


def load_json(path: str) -> Any:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def _is_container_chapter(title: str, description: str) -> bool:
    """Return True only for meta/administrative chapters that should not appear as draft sections."""
    t = title.strip()
    if re.match(r'^[一二三四五六七八九十]+\s*(内容要求|格式要求|排版要求)', t):
        return True
    return False


def _dedup_chapter_groups(brief: list) -> list:
    """If the schema contains two complete groups both starting with a '1.' chapter,
    keep only the last such group."""
    root_one_idx = [i for i, c in enumerate(brief)
                    if re.match(r'^1[\.\s。]', (c.get("title") or "").strip())]
    if len(root_one_idx) >= 2:
        return brief[root_one_idx[-1]:]
    return brief


def _get_deduped_chapters(content_schema: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Return the deduplicated list of chapter dicts from the schema."""
    chapters = content_schema.get("chapters", {})
    brief = []
    for _, v in chapters.items():
        title = (v.get("title") or "").strip()
        desc = (v.get("description") or "").strip()
        if not title or _is_container_chapter(title, desc):
            continue
        brief.append({
            "title": title,
            "description": desc or None,
            "minWordCount": v.get("minWordCount"),
            "required": v.get("required"),
        })
    return _dedup_chapter_groups(brief)


_TERMINOLOGY_PROMPT = {
    "低":  "使用简洁平实的语言，尽量减少专业术语，表达通俗易懂。",
    "中":  "使用适量专业术语，语言规范清晰，符合一般实验报告要求。",
    "高":  "大量使用规范的专业学术术语，表达精准严谨，符合高水平学术论文标准。",
}


def _distribute_chars(
    chapters: List[Dict[str, Any]],
    total: int,
    weights: Optional[Dict[str, float]] = None,
) -> List[int]:
    """Distribute total target chars across chapters using relative weights.

    Each chapter gets max(minWordCount, weight-proportional share, 80) chars.
    """
    n = len(chapters)
    if n == 0:
        return []

    min_counts = []
    for c in chapters:
        try:
            v = float(c.get("minWordCount") or 0)
        except Exception:
            v = 0.0
        min_counts.append(int(v))

    # Relative weights: default 1.0, use provided weights dict if given
    w_values = []
    for c in chapters:
        title = c.get("title", "")
        w = 1.0
        if weights:
            w = max(0.1, float(weights.get(title, 1.0)))
        w_values.append(w)

    total_weight = sum(w_values)
    result = []
    for w, mc in zip(w_values, min_counts):
        allocated = int(total * w / total_weight)
        result.append(max(allocated, mc, 80))
    return result


def _generate_chapter_content(
    client: DeepSeekClient,
    chapter: Dict[str, Any],
    experiment_text: str,
    target_chars: int,
    terminology: str = "中",
) -> str:
    """Call LLM once to generate content for a single chapter. Returns plain text."""
    min_wc = chapter.get("minWordCount")
    try:
        min_wc_int = int(float(min_wc or 0))
    except Exception:
        min_wc_int = 0

    terminology_note = _TERMINOLOGY_PROMPT.get(terminology, _TERMINOLOGY_PROMPT["中"])

    prompt = {
        "instruction": (
            "根据章节要求与实验内容，为以下章节生成正文内容。"
            "只输出纯文本正文，不要输出标题，不要输出JSON，不要添加额外说明。"
            f"目标字数约 {target_chars} 字，最少 {max(min_wc_int, target_chars // 2)} 字。"
            f"语言风格要求：{terminology_note}"
        ),
        "chapter_title": chapter["title"],
        "chapter_description": chapter.get("description") or "",
        "experiment_content": experiment_text,
    }
    messages = [
        {"role": "system", "content": "你是严谨的实验报告写作助手。只输出章节正文内容，不要包含标题，不要输出JSON。"},
        {"role": "user", "content": json.dumps(prompt, ensure_ascii=False)},
    ]
    max_tokens = min(int(target_chars * 1.5) + 300, 4096)
    return client.chat(messages, temperature=0.25, max_tokens=max_tokens).strip()


def generate_draft_chapters(
    content_schema: Dict[str, Any],
    experiment_text: str,
    client: DeepSeekClient,
    total_target: int,
    progress_cb: Optional[Callable[[int, int, str], None]] = None,
    chapter_weights: Optional[Dict[str, float]] = None,
    terminology: str = "中",
) -> List[Dict[str, Any]]:
    """Generate all draft chapters by calling the LLM once per chapter.

    Args:
        progress_cb: optional callback(current_index, total, chapter_title)
    Returns:
        List of {"title": ..., "content": ...}
    """
    chapters = _get_deduped_chapters(content_schema)
    char_budgets = _distribute_chars(chapters, total_target, weights=chapter_weights)
    results = []

    for i, (chapter, budget) in enumerate(zip(chapters, char_budgets)):
        title = chapter["title"]
        if progress_cb:
            progress_cb(i, len(chapters), title)

        effective_budget = max(budget, 80)

        try:
            content = _generate_chapter_content(
                client, chapter, experiment_text, effective_budget, terminology=terminology
            )
        except Exception as e:
            content = f"（生成失败：{e}）"

        results.append({"title": title, "content": content})

    return results


# ── Legacy single-call helpers (kept for backward compat) ────────────────────

def build_prompt(content_schema: Dict[str, Any], experiment_text: str, target_len: int) -> str:
    brief = _get_deduped_chapters(content_schema)
    prompt = {
        "instruction": (
            "根据章节约束与实验内容，为下列每一个章节生成正文内容。"
            "必须为 chapters 列表中的每一项都输出对应内容，不能跳过任何章节。"
            "以严格 JSON 输出，顶层为 chapters 数组，每项包含 title（与输入完全一致）和 content（纯文本正文）。"
            "内容较少的章节可简短，但不能省略。只输出文字，不要插入图片。"
        ),
        "chapters": brief,
        "target_length_chars": target_len,
        "experiment_content": experiment_text,
        "rules": [
            "中文撰写",
            "必须覆盖 chapters 列表中的全部章节，一个都不能少",
            "章节 title 与输入完全一致",
            "满足 minWordCount 要求",
        ],
    }
    return json.dumps(prompt, ensure_ascii=False)


def _parse_llm_response(resp: str) -> List[Dict[str, Any]]:
    s = resp.strip()
    i = s.find("{")
    j = s.rfind("}")
    if i != -1 and j != -1 and j >= i:
        s = s[i : j + 1]
    try:
        obj = json.loads(s)
    except Exception:
        return []
    if isinstance(obj, dict):
        chapters = obj.get("chapters")
        if isinstance(chapters, list):
            return chapters
        result = []
        for k, v in obj.items():
            if isinstance(v, dict):
                result.append({"title": v.get("title") or k, "content": v.get("content") or ""})
        return result
    return []


def _infer_heading_level(title: str) -> int:
    m = re.match(r'^\d+(\.\d+)*', title.strip())
    if not m:
        return 1
    dots = m.group(0).count(".")
    return min(dots + 1, 4)


def write_draft_docx(chapters: List[Dict[str, Any]], out):
    """Write chapters to a docx. `out` may be a file path string or a file-like object."""
    doc = Document()
    for item in chapters:
        title = (item.get("title") or "").strip()
        content = (item.get("content") or "").strip()
        if title:
            level = _infer_heading_level(title)
            doc.add_heading(title, level=level)
        if content:
            for line in content.splitlines():
                line = line.strip()
                if line:
                    doc.add_paragraph(line)
    doc.save(out)


def main():
    cfg = get_env_config()

    ap = argparse.ArgumentParser(prog="gen-draft", description="根据内容schema与实验内容JSON生成纯文字Word草稿")
    ap.add_argument("--content", required=True, help="content schema json 路径")
    ap.add_argument("--experiment", required=True, help="视觉模型提取的实验内容 json 路径")
    ap.add_argument("--out-docx", required=True, help="输出草稿 docx 路径")
    ap.add_argument("--target-length", type=int,
                    default=int(cfg.get("REPORT_TARGET_CHARS") or 3000),
                    help="目标生成字数（默认读取 .env REPORT_TARGET_CHARS）")
    ap.add_argument("--api_key")
    ap.add_argument("--model")
    ap.add_argument("--base_url")
    args = ap.parse_args()

    content_schema = load_json(args.content)
    exp_obj = load_json(args.experiment)
    experiment_text = json.dumps(exp_obj, ensure_ascii=False)

    client = DeepSeekClient(
        api_key=args.api_key or cfg.get("DASHSCOPE_API_KEY"),
        base_url=args.base_url or cfg.get("DASHSCOPE_BASE_URL"),
        model=args.model or cfg.get("DASHSCOPE_MODEL") or "qwen-plus",
    )

    n_chapters = len(_get_deduped_chapters(content_schema))
    print(f"逐章节生成，共 {n_chapters} 个章节，目标总字数 {args.target_length}...")

    def _progress(i, total, title):
        print(f"  [{i+1}/{total}] {title}")

    chapters = generate_draft_chapters(
        content_schema, experiment_text, client,
        total_target=args.target_length,
        progress_cb=_progress,
    )

    os.makedirs(os.path.dirname(args.out_docx), exist_ok=True) if os.path.dirname(args.out_docx) else None
    write_draft_docx(chapters, args.out_docx)

    total_chars = sum(len(c.get("content", "")) for c in chapters)
    print(f"完成：{args.out_docx}  实际字数约 {total_chars}")


if __name__ == "__main__":
    main()
