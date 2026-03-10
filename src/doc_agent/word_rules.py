import re
from typing import Dict, Any, List, Tuple
from docx import Document
from docx.oxml.ns import qn


# ── Heading level detection ───────────────────────────────────────────────────

def _heading_level(style_name: str) -> int:
    s = style_name.strip().lower()
    if s.startswith("heading"):
        try:
            return int(s.replace("heading", "").strip())
        except Exception:
            return 0
    if s.startswith("标题"):
        try:
            return int(s.replace("标题", "").strip())
        except Exception:
            return 0
    return 0


# ── Heading text cleaning ─────────────────────────────────────────────────────

_FORMAT_TOKENS = ["宋体", "黑体", "楷体", "Times", "Arial", "行距", "段前", "段后",
                  "缩进", "居中", "对齐", "加粗", "小四", "四号", "五号", "Heading",
                  "样式", "不分页", "同页"]


def _clean_heading_text(text: str) -> str:
    """Extract the real chapter title from a heading that may embed format specs.

    Handles patterns like:
      '正文 ((宋体，小四号...))\n2.2 实验方法'
      '1. 实验原理（样式:标题1；黑体...）'
    """
    if not text:
        return text
    lines = [l.strip() for l in text.splitlines()]
    real_lines = []
    for line in lines:
        # Skip lines that are purely format spec lines
        if "正文 ((" in line or "正文((" in line:
            continue
        if "((" in line and "))" in line:
            # Strip the ((...)) block
            line = re.sub(r'\(\([^)]*\)\)', '', line).strip()
        # Strip inline format annotations like （样式:...）
        line = re.sub(r'（[^）]{0,60}）', '', line).strip()
        if not line:
            continue
        # Drop lines consisting only of format keywords
        if all(tok in line for tok in []) or any(
            all(tok in line for tok in pair)
            for pair in [("宋体",), ("黑体",), ("行距",), ("缩进",)]
        ):
            pass
        real_lines.append(line)
    return real_lines[0] if real_lines else ""


# ── Page / section info ───────────────────────────────────────────────────────

def _get_sections_info(doc: Document) -> Dict[str, Any]:
    secs = []
    for s in doc.sections:
        secs.append({
            "page_width": int(s.page_width),
            "page_height": int(s.page_height),
            "orientation": "landscape" if s.orientation == 1 else "portrait",
            "margins": {
                "top": int(s.top_margin),
                "bottom": int(s.bottom_margin),
                "left": int(s.left_margin),
                "right": int(s.right_margin),
                "header": int(s.header_distance),
                "footer": int(s.footer_distance),
                "gutter": int(s.gutter),
            },
        })
    return {"sections": secs}


# ── Full text ─────────────────────────────────────────────────────────────────

def _collect_text(doc: Document) -> str:
    parts: List[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for tbl in doc.tables:
        for r in tbl.rows:
            row_text = []
            for c in r.cells:
                tx = c.text.strip()
                if tx:
                    row_text.append(tx)
            if row_text:
                parts.append(" | ".join(row_text))
    return "\n".join(parts)


# ── Headings ──────────────────────────────────────────────────────────────────

def _get_headings(doc: Document) -> List[Dict[str, Any]]:
    res = []
    for p in doc.paragraphs:
        if p.style and p.style.name:
            lvl = _heading_level(p.style.name)
            if lvl > 0:
                raw = p.text.strip()
                clean = _clean_heading_text(raw)
                if clean:
                    res.append({"text": clean, "level": lvl, "style": p.style.name})
    return res


def _get_heading_blocks(doc: Document) -> List[Dict[str, Any]]:
    blocks: List[Dict[str, Any]] = []
    current = None
    for p in doc.paragraphs:
        name = p.style.name if p.style else ""
        lvl = _heading_level(name) if name else 0
        if lvl > 0:
            if current:
                blocks.append(current)
            clean = _clean_heading_text(p.text.strip())
            if not clean:
                current = None
                continue
            current = {"text": clean, "level": lvl, "style": name, "body": []}
        else:
            if current is not None:
                t = p.text.strip()
                if t:
                    current["body"].append(t)
    if current:
        blocks.append(current)
    for b in blocks:
        b["content"] = "\n".join(b.get("body", []))
        if "body" in b:
            del b["body"]
    return blocks


# ── Paragraph styles (paragraph-level overrides) ──────────────────────────────

def _get_paragraph_styles(doc: Document) -> List[Dict[str, Any]]:
    acc: Dict[str, Dict[str, Any]] = {}
    for p in doc.paragraphs:
        name = p.style.name if p.style else "Normal"
        fmt = p.paragraph_format
        entry = acc.get(name) or {
            "style": name,
            "alignment": str(fmt.alignment) if fmt.alignment is not None else None,
            "line_spacing": float(fmt.line_spacing) if fmt.line_spacing else None,
            "space_before": int(fmt.space_before) if fmt.space_before else None,
            "space_after": int(fmt.space_after) if fmt.space_after else None,
            "first_line_indent": int(fmt.first_line_indent) if fmt.first_line_indent else None,
            "left_indent": int(fmt.left_indent) if fmt.left_indent else None,
            "right_indent": int(fmt.right_indent) if fmt.right_indent else None,
            "count": 0,
        }
        entry["count"] = entry["count"] + 1
        acc[name] = entry
    return list(acc.values())


# ── Style definitions from XML ────────────────────────────────────────────────

def _get_style_definitions_xml(doc: Document) -> Dict[str, Dict[str, Any]]:
    """Extract font, size, bold, italic, line-spacing from the docx style XML.

    python-docx's paragraph_format only reflects paragraph-level overrides.
    Style definitions (the actual inherited values) live in the XML and must
    be read directly.
    """
    result: Dict[str, Dict[str, Any]] = {}
    for style in doc.styles:
        try:
            elem = style.element
            if elem is None:
                continue
            name = style.name
            info: Dict[str, Any] = {}

            # ── Run properties: font, size, bold, italic ──────────────────────
            rpr = elem.find('.//' + qn('w:rPr'))
            if rpr is not None:
                rfonts = rpr.find(qn('w:rFonts'))
                if rfonts is not None:
                    ea = rfonts.get(qn('w:eastAsia'))
                    ascii_f = rfonts.get(qn('w:ascii'))
                    if ea:
                        info['fontEastAsia'] = ea
                    if ascii_f:
                        info['fontAscii'] = ascii_f
                sz = rpr.find(qn('w:sz'))
                if sz is not None:
                    val = sz.get(qn('w:val'))
                    if val:
                        try:
                            info['fontSizePt'] = int(val) / 2.0
                        except Exception:
                            pass
                b = rpr.find(qn('w:b'))
                if b is not None:
                    bval = b.get(qn('w:val'))
                    info['bold'] = bval is None or bval.lower() not in ('false', '0', 'off')
                i_elem = rpr.find(qn('w:i'))
                if i_elem is not None:
                    ival = i_elem.get(qn('w:val'))
                    info['italic'] = ival is None or ival.lower() not in ('false', '0', 'off')

            # ── Paragraph properties: spacing, indent, alignment ──────────────
            ppr = elem.find('.//' + qn('w:pPr'))
            if ppr is not None:
                spacing = ppr.find(qn('w:spacing'))
                if spacing is not None:
                    line = spacing.get(qn('w:line'))
                    line_rule = spacing.get(qn('w:lineRule'))
                    before = spacing.get(qn('w:before'))
                    after = spacing.get(qn('w:after'))
                    if line:
                        try:
                            line_val = int(line)
                            if line_rule in ('auto', None):
                                info['lineSpacing'] = round(line_val / 240.0, 4)
                            else:
                                info['lineSpacingPt'] = round(line_val / 20.0, 2)
                        except Exception:
                            pass
                    if before:
                        try:
                            info['spaceBeforePt'] = round(int(before) / 20.0, 2)
                        except Exception:
                            pass
                    if after:
                        try:
                            info['spaceAfterPt'] = round(int(after) / 20.0, 2)
                        except Exception:
                            pass
                ind = ppr.find(qn('w:ind'))
                if ind is not None:
                    first = ind.get(qn('w:firstLine'))
                    left = ind.get(qn('w:left'))
                    right = ind.get(qn('w:right'))
                    # Convert twips → EMU (1 twip = 635 EMU)
                    if first:
                        try:
                            info['firstLineIndentEmu'] = int(first) * 635
                        except Exception:
                            pass
                    if left:
                        try:
                            info['leftIndentEmu'] = int(left) * 635
                        except Exception:
                            pass
                    if right:
                        try:
                            info['rightIndentEmu'] = int(right) * 635
                        except Exception:
                            pass
                jc = ppr.find(qn('w:jc'))
                if jc is not None:
                    info['alignment'] = jc.get(qn('w:val'))

            if info:
                result[name] = info
        except Exception:
            pass
    return result


# ── Table styles ──────────────────────────────────────────────────────────────

def _get_table_styles(doc: Document) -> Dict[str, Any]:
    styles: Dict[str, int] = {}
    for t in doc.tables:
        n = t.style.name if t.style else "Table Grid"
        styles[n] = styles.get(n, 0) + 1
    return {"table_styles": [{"style": k, "count": v} for k, v in styles.items()]}


# ── Caption detection ─────────────────────────────────────────────────────────

def _detect_captions(doc: Document) -> Dict[str, Any]:
    keys = ["caption", "图题", "表题", "图名", "表名"]
    caps: List[Dict[str, Any]] = []
    for p in doc.paragraphs:
        if p.style and p.style.name:
            sname = p.style.name.lower()
            if any(k in sname for k in keys):
                tx = p.text.strip()
                has_num = ("图" in tx and any(ch.isdigit() for ch in tx)) or (
                    "表" in tx and any(ch.isdigit() for ch in tx)
                )
                caps.append({"style": p.style.name, "text": tx, "has_numbering": has_num})
    return {"captions": caps}


# ── Content pattern detection ─────────────────────────────────────────────────

def _guess_required_sections(headings: List[Dict[str, Any]], full_text: str) -> List[str]:
    known = ["摘要", "引言", "实验目的", "实验原理", "实验器材", "实验设备",
             "实验步骤", "实验数据", "结果与分析", "结果分析", "讨论", "结论",
             "参考文献", "附录"]
    found = []
    names = [h["text"] for h in headings]
    for k in known:
        if any(k in n for n in names):
            found.append(k)
    dedup = []
    for x in found:
        if x not in dedup:
            dedup.append(x)
    return dedup


def _content_patterns(full_text: str) -> Dict[str, Any]:
    rules = {"data_recording": [], "analysis_methods": [], "conclusion_requirements": [], "placeholders": []}
    for line in full_text.splitlines():
        t = line.strip()
        if not t:
            continue
        if any(k in t for k in ["单位", "小数", "保留", "测量", "记录", "表格"]):
            rules["data_recording"].append(t)
        if any(k in t for k in ["方法", "计算", "拟合", "误差", "分析", "处理"]):
            rules["analysis_methods"].append(t)
        if any(k in t for k in ["结论", "得到", "说明", "验证"]):
            rules["conclusion_requirements"].append(t)
        if any(k in t for k in ["{{", "}}", "请在此", "填写", "须"]):
            rules["placeholders"].append(t)
    for k in list(rules.keys()):
        uniq = []
        for v in rules[k]:
            if v not in uniq:
                uniq.append(v)
        rules[k] = uniq[:50]
    return rules


# ── Body style extraction from format-spec annotations ───────────────────────

_FONT_NAMES = ["宋体", "黑体", "楷体", "仿宋", "Times New Roman", "Arial", "Calibri"]
_SIZE_MAP = {"小四": 12.0, "四号": 14.0, "小三": 15.0, "三号": 16.0, "五号": 10.5, "小五": 9.0}
_SPEC_RE = re.compile(r'\(\(([^)]+)\){1,2}')


def _extract_body_style_from_annotations(doc: Document) -> Dict[str, Any]:
    """Parse '((宋体，小四号，1.5倍行距...))' annotations in heading texts to get
    the body paragraph format spec. These annotations are written by template
    authors to describe the intended Normal/body style, which Word itself stores
    via theme fonts (no explicit w:rFonts on the Normal style XML)."""
    body: Dict[str, Any] = {}
    for p in doc.paragraphs:
        if not (p.style and p.style.name):
            continue
        if _heading_level(p.style.name) == 0:
            continue
        text = p.text
        if "正文" not in text:
            continue
        m = _SPEC_RE.search(text)
        if not m:
            continue
        spec = m.group(1)
        if not body.get("fontEastAsia"):
            for fn in _FONT_NAMES:
                if fn in spec:
                    body["fontEastAsia"] = fn
                    break
        if not body.get("fontSizePt"):
            for sn, sp in _SIZE_MAP.items():
                if sn in spec:
                    body["fontSizePt"] = sp
                    break
        if not body.get("lineSpacing") and "1.5" in spec:
            body["lineSpacing"] = 1.5
        if body:
            break
    return body


# ── Main entry ────────────────────────────────────────────────────────────────

def analyze_docx(path: str) -> Dict[str, Any]:
    doc = Document(path)
    full_text = _collect_text(doc)
    headings = _get_headings(doc)
    heading_blocks = _get_heading_blocks(doc)
    style_defs = _get_style_definitions_xml(doc)

    # Supplement Normal style with info from template format annotations
    body_from_text = _extract_body_style_from_annotations(doc)
    if body_from_text:
        normal = style_defs.get("Normal", {})
        for k, v in body_from_text.items():
            if not normal.get(k):
                normal[k] = v
        style_defs["Normal"] = normal

    format_spec: Dict[str, Any] = {}
    format_spec.update(_get_sections_info(doc))
    format_spec["paragraph_styles"] = _get_paragraph_styles(doc)
    format_spec["style_definitions"] = style_defs
    format_spec.update(_get_table_styles(doc))
    format_spec.update(_detect_captions(doc))
    content_spec: Dict[str, Any] = {}
    content_spec["required_sections"] = _guess_required_sections(headings, full_text)
    content_spec.update(_content_patterns(full_text))
    content_spec["heading_outline"] = headings
    content_spec["heading_blocks"] = heading_blocks
    return {
        "format_spec": format_spec,
        "content_spec": content_spec,
        "full_text_sample": "\n".join(full_text.splitlines()[:50]),
    }
