import json
import shutil
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _emu_to_inches(v):
    try:
        return float(v) / 914400.0
    except Exception:
        return None


def render_report(format_schema: Dict[str, Any], generated: Dict[str, Any], out_path: str):
    doc = Document()
    docs = format_schema.get("documents", [])
    fm = docs[0] if docs else {}
    page = fm.get("pageLayout", {})
    sec = doc.sections[0]
    ori = page.get("orientation")
    if ori == "landscape":
        sec.orientation = WD_ORIENT.LANDSCAPE
    w = page.get("sizeEmu", {}).get("width")
    h = page.get("sizeEmu", {}).get("height")
    wi = _emu_to_inches(w)
    hi = _emu_to_inches(h)
    if wi and hi:
        sec.page_width, sec.page_height = Inches(wi), Inches(hi)
    m = page.get("marginsEmu", {})
    for k in ["top", "bottom", "left", "right"]:
        val = m.get(k)
        iv = _emu_to_inches(val)
        if iv:
            setattr(sec, f"{k}_margin", Inches(iv))
    headings = fm.get("headingLevels", [])
    level_map = {}
    for hdef in headings:
        t = hdef.get("name")
        lvl = hdef.get("level")
        try:
            lvl_f = int(float(lvl))
        except Exception:
            lvl_f = 1
        if isinstance(t, str):
            level_map[t] = lvl_f
    chapters = generated
    if "chapters" in generated:
        chapters = generated.get("chapters", {})
    order: List[str] = []
    for hdef in headings:
        nm = hdef.get("name")
        if nm in chapters:
            order.append(nm)
    for k in chapters.keys():
        if k not in order:
            order.append(k)
    for key in order:
        item = chapters.get(key) or {}
        title = item.get("title") or key
        content = item.get("content") or ""
        lv = level_map.get(title, 1)
        try:
            p = doc.add_paragraph(title, style=f"Heading {lv}")
        except Exception:
            p = doc.add_paragraph(title)
        for line in content.splitlines():
            doc.add_paragraph(line)
    doc.save(out_path)


def read_draft_docx(path: str) -> Dict[str, Any]:
    """Read a draft Word document and extract chapters by heading style.

    Returns a dict compatible with what generator.py produces:
    {"chapters": {"key": {"title": ..., "content": ...}}}
    """
    doc = Document(path)
    chapters: Dict[str, Any] = {}
    order: List[str] = []
    current_key: str = None
    current_title: str = None
    current_lines: List[str] = []

    def _flush():
        if current_key is not None:
            chapters[current_key] = {
                "title": current_title,
                "content": "\n".join(current_lines),
            }

    for p in doc.paragraphs:
        style_name = p.style.name if p.style else "Normal"
        is_heading = style_name.lower().startswith("heading") or style_name.startswith("标题")
        if is_heading:
            _flush()
            current_title = p.text.strip()
            current_key = current_title.replace(" ", "") or f"section_{len(chapters)}"
            order.append(current_key)
            current_lines = []
        else:
            text = p.text.strip()
            if text:
                current_lines.append(text)

    _flush()
    return {"chapters": chapters, "order": order}


def _apply_page_layout(doc: Document, fm: Dict[str, Any]):
    page = fm.get("pageLayout", {})
    sec = doc.sections[0]
    ori = page.get("orientation")
    if ori == "landscape":
        sec.orientation = WD_ORIENT.LANDSCAPE
    w = page.get("sizeEmu", {}).get("width")
    h = page.get("sizeEmu", {}).get("height")
    wi = _emu_to_inches(w)
    hi = _emu_to_inches(h)
    if wi and hi:
        sec.page_width, sec.page_height = Inches(wi), Inches(hi)
    m = page.get("marginsEmu", {})
    for k in ["top", "bottom", "left", "right"]:
        val = m.get(k)
        iv = _emu_to_inches(val)
        if iv:
            setattr(sec, f"{k}_margin", Inches(iv))


def _apply_para_style(p, spec: Dict[str, Any]):
    """Apply font/spacing properties from a paragraphStyle spec to a paragraph."""
    fmt = p.paragraph_format
    line_spacing = spec.get("lineSpacing")
    if line_spacing is not None:
        try:
            fmt.line_spacing = float(line_spacing)
        except Exception:
            pass
    indent = spec.get("indent", {}) or {}
    first = indent.get("firstLineEmu")
    left = indent.get("leftEmu")
    if first:
        iv = _emu_to_inches(first)
        if iv:
            fmt.first_line_indent = Inches(iv)
    if left:
        iv = _emu_to_inches(left)
        if iv:
            fmt.left_indent = Inches(iv)
    font_family = spec.get("fontFamily")
    font_size = spec.get("fontSizePt")
    bold = spec.get("bold")
    italic = spec.get("italic")
    if any(v is not None for v in [font_family, font_size, bold, italic]):
        for run in p.runs:
            if font_family:
                run.font.name = font_family
                # Also set East Asian font via XML to handle Chinese characters
                try:
                    from docx.oxml.ns import qn
                    from lxml import etree
                    rpr = run._r.get_or_add_rPr()
                    rfonts = rpr.find(qn("w:rFonts"))
                    if rfonts is None:
                        rfonts = etree.SubElement(rpr, qn("w:rFonts"))
                    rfonts.set(qn("w:eastAsia"), font_family)
                except Exception:
                    pass
            if font_size is not None:
                try:
                    run.font.size = Pt(float(font_size))
                except Exception:
                    pass
            if bold is not None:
                run.font.bold = bold
            if italic is not None:
                run.font.italic = italic


def _para_has_image(para) -> bool:
    """Return True if this paragraph contains at least one inline image (w:drawing)."""
    try:
        from docx.oxml.ns import qn
        return para._p.find('.//' + qn('w:drawing')) is not None
    except Exception:
        return False


def _collect_inline_shapes(doc: Document):
    """Yield all inline shapes from paragraphs AND table cells."""
    from docx.oxml.ns import qn
    from docx.shape import InlineShape

    def _shapes_in_para(p):
        for drawing in p._p.findall('.//' + qn('w:drawing')):
            inline = drawing.find('.//' + qn('wp:inline'))
            if inline is not None:
                try:
                    yield InlineShape(inline)
                except Exception:
                    pass

    for p in doc.paragraphs:
        yield from _shapes_in_para(p)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield from _shapes_in_para(p)


def _resize_images(doc: Document, max_width_emu: int, max_height_emu: int):
    """Scale inline images to fit within max_width × max_height, preserving aspect ratio.
    Covers both body paragraphs and table cells.
    """
    for shape in _collect_inline_shapes(doc):
        try:
            w = shape.width
            h = shape.height
            if not w or not h:
                continue
            scale = 1.0
            if w > max_width_emu:
                scale = min(scale, max_width_emu / w)
            if h * scale > max_height_emu:
                scale = min(scale, max_height_emu / h)
            if scale < 1.0:
                shape.width = int(w * scale)
                shape.height = int(h * scale)
        except Exception:
            pass


def _format_image_para(para):
    """Center-align an image paragraph and clear indent/font overrides that cause misalignment."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    try:
        fmt = para.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.first_line_indent = None
        fmt.left_indent = None
        fmt.right_indent = None
        fmt.space_before = Pt(6)
        fmt.space_after = Pt(6)
        # Remove any run-level font overrides so they don't affect image anchor
        for run in para.runs:
            if not run.text.strip():   # runs with only whitespace / no text
                run.font.name = None
                run.font.size = None
    except Exception:
        pass


def _insert_student_info(doc: Document, info: Dict[str, str]):
    """Insert a student info block at the very beginning of the document."""
    fields = [
        ("姓    名", info.get("name", "")),
        ("学    号", info.get("student_id", "")),
        ("课程名称", info.get("course", "")),
        ("指导教师", info.get("teacher", "")),
        ("实验日期", info.get("date", "")),
    ]
    # Only insert if at least one field is filled
    if not any(v for _, v in fields):
        return

    # Insert paragraphs before the first element (reverse order using XML addprevious)
    first_p = doc.paragraphs[0]._p if doc.paragraphs else None

    def _make_info_para(label: str, value: str) -> Any:
        from lxml import etree
        from docx.oxml.ns import qn
        p = etree.SubElement(etree.Element("root"), qn("w:p"))
        # Build paragraph via python-docx API on a temp doc, then steal the element
        tmp = doc.add_paragraph()
        tmp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tmp.paragraph_format.space_before = Pt(2)
        tmp.paragraph_format.space_after = Pt(2)
        run = tmp.add_run(f"{label}：{value}")
        run.font.size = Pt(12)
        return tmp

    # Add a blank line spacer at top, then fields, then a divider line
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)

    info_paras = []
    for label, value in fields:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(f"{label}：{value if value else '_______________'}")
        run.font.size = Pt(11)
        info_paras.append(p)

    divider = doc.add_paragraph("─" * 30)
    divider.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider.paragraph_format.space_after = Pt(8)
    for run in divider.runs:
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # Move all inserted paragraphs before the original first paragraph
    if first_p is not None:
        for p in [spacer, *info_paras, divider]:
            first_p.addprevious(p._p)
        # Remove the now-appended copies at end
        body = doc.element.body
        for p in [spacer, *info_paras, divider]:
            if p._p in body:
                body.remove(p._p)
        # Re-insert in order
        for p in reversed([spacer, *info_paras, divider]):
            first_p.addprevious(p._p)


def _add_figure_captions(doc: Document):
    """Insert 'Figure N' caption after each image paragraph (skip if next para already looks like a caption)."""
    from docx.oxml.ns import qn
    from lxml import etree

    # Collect indices of image paragraphs first
    paras = list(doc.paragraphs)
    image_indices = [i for i, p in enumerate(paras) if _para_has_image(p)]

    # Forward pass: assign figure numbers in document order, skipping existing captions
    fig_num = 1
    insert_plan = []  # list of (idx, fig_num) for indices that need a caption
    for idx in image_indices:
        next_text = paras[idx + 1].text.strip() if idx + 1 < len(paras) else ""
        if next_text.startswith("图") or next_text.startswith("Fig"):
            fig_num += 1
            continue
        insert_plan.append((idx, fig_num))
        fig_num += 1

    # Reverse pass: insert captions so earlier indices are not shifted
    for idx, num in reversed(insert_plan):
        para = paras[idx]
        new_para = doc.add_paragraph()
        new_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        new_para.paragraph_format.space_before = Pt(2)
        new_para.paragraph_format.space_after = Pt(6)
        run = new_para.add_run(f"图{num}  ")
        run.font.size = Pt(10.5)

        # Move the new paragraph element to right after the image paragraph
        para._p.addnext(new_para._p)


def render_report_from_draft(
    draft_path: str,
    format_schema: Dict[str, Any],
    out_path: str,
    student_info: Optional[Dict[str, str]] = None,
    add_figure_captions: bool = False,
):
    """Apply format_schema page layout and styles to an existing draft docx, save to out_path.

    Preserves all user edits including embedded images and tables.
    Images are scaled to fit within the usable page area and centered.
    """
    shutil.copy2(draft_path, out_path)
    doc = Document(out_path)
    docs = format_schema.get("documents", [])
    fm = docs[0] if docs else {}

    _apply_page_layout(doc, fm)

    # Calculate usable area from format_schema for image sizing
    page = fm.get("pageLayout", {})
    margins = page.get("marginsEmu", {})
    page_w = _emu_to_inches(page.get("sizeEmu", {}).get("width")) or 8.27
    left_m = _emu_to_inches(margins.get("left")) or 1.0
    right_m = _emu_to_inches(margins.get("right")) or 1.0
    top_m = _emu_to_inches(margins.get("top")) or 1.0
    bottom_m = _emu_to_inches(margins.get("bottom")) or 1.0
    page_h = _emu_to_inches(page.get("sizeEmu", {}).get("height")) or 11.69

    usable_w_emu = int((page_w - left_m - right_m) * 914400)
    # Limit image height to 45% of usable page height to keep surrounding context visible
    usable_h_emu = int((page_h - top_m - bottom_m) * 914400 * 0.45)

    _resize_images(doc, usable_w_emu, usable_h_emu)

    if add_figure_captions:
        _add_figure_captions(doc)

    if student_info:
        _insert_student_info(doc, student_info)

    # Build style lookup: name -> spec
    para_styles = fm.get("paragraphStyles", [])
    style_map: Dict[str, Dict[str, Any]] = {s.get("name"): s for s in para_styles if s.get("name")}

    # Pandoc uses "Body Text"/"First Paragraph" etc. instead of "Normal";
    # map them to the Normal spec so body paragraphs get formatted correctly.
    _NORMAL_ALIASES = {"Body Text", "First Paragraph", "Compact", "Body Text First Indent"}
    normal_spec = style_map.get("Normal")
    if normal_spec:
        for alias in _NORMAL_ALIASES:
            if alias not in style_map:
                style_map[alias] = normal_spec

    for p in doc.paragraphs:
        if _para_has_image(p):
            # Image paragraph: center + clean indents; do NOT apply body text style
            _format_image_para(p)
        else:
            style_name = p.style.name if p.style else "Normal"
            spec = style_map.get(style_name)
            if spec:
                _apply_para_style(p, spec)

    doc.save(out_path)
