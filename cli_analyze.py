import argparse
import json
import os
import re
import glob
from src.doc_agent.word_rules import analyze_docx
from src.doc_agent.classify import refine_with_llm
from src.doc_agent.exporter import export_markdown, export_docx
from src.img_agent.deepseek_client import DeepSeekClient
from src.doc_agent.normalize import normalize_spec


def run_analyze(template_path: str, out_json: str, out_md: str = None, out_docx: str = None, use_llm: bool = True, api_key: str = None, model: str = None, base_url: str = None):
    base = analyze_docx(template_path)
    res = base
    if use_llm:
        client = DeepSeekClient(api_key=api_key, model=model or "qwen-plus", base_url=base_url)
        llm_obj = refine_with_llm(base, client)
        res = normalize_spec(base, llm_obj, template_path)
    else:
        res = normalize_spec(base, {}, template_path)
    os.makedirs(os.path.dirname(out_json), exist_ok=True) if os.path.dirname(out_json) else None
    with open(out_json, "w", encoding="utf-8") as f:
        f.write(json.dumps(res, ensure_ascii=False, indent=2))
    if out_md:
        md = export_markdown(res)
        os.makedirs(os.path.dirname(out_md), exist_ok=True) if os.path.dirname(out_md) else None
        with open(out_md, "w", encoding="utf-8") as f:
            f.write(md)
    if out_docx:
        export_docx(res, out_docx)
    print(out_json)
    if out_md:
        print(out_md)
    if out_docx:
        print(out_docx)


def run_demo(out_dir: str, api_key: str = None, model: str = None, base_url: str = None):
    from docx import Document
    os.makedirs(out_dir, exist_ok=True)
    p = os.path.join(out_dir, "demo_template.docx")
    d = Document()
    d.add_paragraph("实验报告模板", style=None)
    d.add_paragraph("摘要", style="Heading 1")
    d.add_paragraph("请在此填写摘要。建议200字以内。")
    d.add_paragraph("实验目的", style="Heading 1")
    d.add_paragraph("应明确列出实验目标。")
    d.add_paragraph("实验原理", style="Heading 1")
    d.add_paragraph("应阐述基本原理、公式与符号，计算方法。")
    d.add_paragraph("实验步骤", style="Heading 1")
    d.add_paragraph("步骤应编号，数据记录表需包含单位与小数位。")
    d.add_paragraph("实验数据", style="Heading 1")
    d.add_paragraph("建议以表格形式记录测量数据，注明单位。")
    d.add_paragraph("结果与分析", style="Heading 1")
    d.add_paragraph("需给出分析方法、误差评估与拟合。")
    d.add_paragraph("结论", style="Heading 1")
    d.add_paragraph("应简洁准确，呼应目的。")
    d.add_paragraph("参考文献", style="Heading 1")
    d.add_paragraph("按规范格式书写。")
    d.save(p)
    out_json = os.path.join(out_dir, "demo_rules.json")
    out_md = os.path.join(out_dir, "demo_rules.md")
    out_docx = os.path.join(out_dir, "demo_rules.docx")
    run_analyze(p, out_json, out_md, out_docx, use_llm=True, api_key=api_key, model=model, base_url=base_url)


def main():
    ap = argparse.ArgumentParser(prog="tpl-analyze", description="实验报告模板规范解析智能体")
    sub = ap.add_subparsers(dest="cmd")
    p1 = sub.add_parser("analyze")
    group = p1.add_mutually_exclusive_group(required=True)
    group.add_argument("--template")
    group.add_argument("--dir")
    p1.add_argument("--out-json", required=True, help="若使用 --dir，此参数作为输出目录")
    p1.add_argument("--out-md")
    p1.add_argument("--out-docx")
    p1.add_argument("--no-llm", action="store_true")
    p1.add_argument("--api_key")
    p1.add_argument("--model", default="qwen-plus")
    p1.add_argument("--base_url", default="https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions")
    p2 = sub.add_parser("demo")
    p2.add_argument("--out-dir", required=True)
    p2.add_argument("--api_key")
    p2.add_argument("--model", default="qwen-plus")
    p2.add_argument("--base_url", default="https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions")
    args = ap.parse_args()
    if args.cmd == "analyze":
        if getattr(args, "dir", None):
            out_dir = args.out_json
            os.makedirs(out_dir, exist_ok=True)
            paths = []
            for pat in ("*.docx", "*.DOCX"):
                paths.extend(glob.glob(os.path.join(args.dir, pat)))
            if not paths:
                print("No .docx found in dir")
                return
            for p in paths:
                base = os.path.splitext(os.path.basename(p))[0]
                safe = re.sub(r"[^0-9A-Za-z\u4e00-\u9fa5]+", "_", base).strip("_")
                json_path = os.path.join(out_dir, f"{safe}_rules.json")
                md_path = os.path.join(out_dir, f"{safe}_rules.md") if args.out_md else None
                docx_path = os.path.join(out_dir, f"{safe}_rules.docx") if args.out_docx else None
                run_analyze(
                    p,
                    json_path,
                    out_md=md_path,
                    out_docx=docx_path,
                    use_llm=not args.no_llm,
                    api_key=args.api_key,
                    model=args.model,
                    base_url=args.base_url,
                )
        else:
            run_analyze(
                args.template,
                args.out_json,
                out_md=args.out_md,
                out_docx=args.out_docx,
                use_llm=not args.no_llm,
                api_key=args.api_key,
                model=args.model,
                base_url=args.base_url,
            )
    elif args.cmd == "demo":
        run_demo(args.out_dir, api_key=args.api_key, model=args.model, base_url=args.base_url)
    else:
        ap.print_help()


if __name__ == "__main__":
    main()
